import PySimpleGUI as sg
import operator
import aqdefreader
import pandas as pd
import xlsxwriter
import numpy as np
from io import BytesIO
from PIL import Image, ImageDraw
import io
import matplotlib.pyplot as plt
import win32com.client
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from matplotlib import pyplot as plt  
from docx.shared import Cm

df = pd.DataFrame()
dfRowNumber = 0
dfColumnNumber = 0
shotNumber = 0
kavNumber = 0
progressPercentage = 0
outputType = 0
choices = ['.xlsx', '.pdf', '.doc']

def trunc(values, decs=0):
    return np.trunc(values*10**decs)/(10**decs)

def update_title(table, newheadings):
    for cid, text in zip(headings, newheadings):
        table.heading(cid, text=text)

# ------ Some functions to help generate data for the table ------
def make_table(num_rows, num_cols):
    data = [[0 for _ in range(num_cols)] for _ in range(num_rows)]
    # data[0] = [0 for _ in range(num_cols)]
    # for i in range(1, num_rows):
    #     data[i] = [0 for _ in range(num_cols)]
    return data

# ------ Make the Table Data ------
data = make_table(3, 6)
headings = ['1 min', 'Shot ', 'Kav ', '1 max', 'Shot', 'Kav']
updatedHeadings = []
dataOther = []

def icon(check):
    box = (18, 18)
    background = (255, 255, 255, 0)
    rectangle = (2, 2, 16, 16)
    line = ((4, 9), (7, 12), (12, 7))
    im = Image.new('RGBA', box, background)
    draw = ImageDraw.Draw(im, 'RGBA')
    draw.rectangle(rectangle, outline='black', width=1)
    if check == 1:
        draw.line(line, fill='black', width=2, joint='curve')
    elif check == 2:
        draw.line(line, fill='grey', width=2, joint='curve')
    with BytesIO() as output:
        im.save(output, format="PNG")
        png = output.getvalue()
    return png

check = [icon(0), icon(1), icon(2)]

headingsTree = ['K2002-Merkmal', 'K2003-Description']

column_headers = ['']
treedata = sg.TreeData()
treeSeq = []

def sort_table(table, cols):
    """ sort a table by multiple columns
        table: a list of lists (or tuple of tuples) where each inner list
               represents a row
        cols:  a list (or tuple) specifying the column numbers to sort by
               e.g. (1,0) would sort by column 1, then by column 0
    """
    for col in reversed(cols):
        try:
            table = sorted(table, key=operator.itemgetter(col))
        except Exception as e:
            sg.popup_error('Error in sort_table', 'Exception in sort_table', e)
    return table

def SetNumber():
    global shotNumber
    global kavNumber
    global data
    col_layout = [[sg.Ok()]]
    form_rows = [[sg.Text('Total number of data'), sg.Text(dfRowNumber)],
                 [sg.Text('Current number of shot'), sg.Text(shotNumber)],
                 [sg.Text('Current number of kav'), sg.Text(kavNumber)],
                 [sg.Text('Enter Number of Shot and Kav')],                 
                 [sg.Text('Shot', size=(15, 1)),sg.InputText(key='-shot-')],
                 [sg.Text('Kav', size=(15, 1)), sg.InputText(key='-kav-')],
                 [sg.Column(col_layout, element_justification='right', expand_x=True)]]

    window = sg.Window('Setting', form_rows, keep_on_top=True)
    event, values = window.read()
    if event == 'Ok':
        if values['-shot-'] and values['-kav-'] and int(values['-shot-']) * int(values['-kav-']) <= dfRowNumber:
            
            shotNumber = int(values['-shot-'])
            kavNumber = int(values['-kav-'])

            for i in range(3):
                data[i] = [data[i][0], 1, i + 1, data[i][3], 1, i + 1]

    window.close()

def processDFQ(filename):
    dqfFile = aqdefreader.read_dfq_file(filename)
    global df
    df = aqdefreader.create_column_dataframe(dqfFile)

    global column_headers
    global dataOther
    global dfColumnNumber
    global dfRowNumber
    global kavNumber
    global shotNumber

    column_headers = list(df.columns.values)
    dfColumnNumber = len(column_headers)
    dfRowNumber = len(df)
    print("dfRowNumber :", dfRowNumber)
    print("dfColumnNumber :", dfColumnNumber)

    data2003 = [f'{i}' for i in range(dfColumnNumber)]
    dataOther = [[0 for _ in range(3)] for _ in range(dfColumnNumber)]

    k = 0

    measurement = dqfFile.parts[0].get_characteristic_by_index(0).get_measurements()[0].as_dictionary()
    lastMeasurement = dqfFile.parts[0].get_characteristic_by_index(0).get_last_measurement().as_dictionary()
    kavNumber = int(measurement['nest_no'])
    if len(measurement['batch_no']) == 1:
        shotNumber = 0
    else:
        shotNumber = int(measurement['batch_no'][1:])
    if shotNumber:
        shotNumber = int(lastMeasurement['batch_no'][1:])
    if kavNumber:
        kavNumber = int(lastMeasurement['nest_no'])

    print("shotNumber :", shotNumber)
    print("kavNumber :", kavNumber)

    for _, characteristic in enumerate(
        dqfFile.parts[0].get_characteristics()
    ):
        data2003[k] = characteristic.get_data("K2003")
        dataOther[k][0] = float(characteristic.get_data("K2101")) if characteristic.get_data("K2101") else None
        dataOther[k][1] = float(characteristic.get_data("K2111")) if characteristic.get_data("K2111") else None
        dataOther[k][2] = float(characteristic.get_data("K2110")) if characteristic.get_data("K2110") else None
        k += 1
        if k == dfColumnNumber:
            break

    numRow = 3
    numCol = 6

    global data
    global treedata

    data = [[j for j in range(numCol)] for i in range(numRow)]
    
    for i in range(numRow):
        if kavNumber and shotNumber:
            data[i] = [df.get(column_headers[0]).iloc[i], 1, i + 1, df.get(column_headers[1]).iloc[i], 1, i + 1]
        elif shotNumber:
            data[i] = [df.get(column_headers[0]).iloc[i], 1, 0, df.get(column_headers[1]).iloc[i], 1, 0]
        elif kavNumber:
            data[i] = [df.get(column_headers[0]).iloc[i], 0, i + 1, df.get(column_headers[1]).iloc[i], 0, i + 1]
        else:
            data[i] = [df.get(column_headers[0]).iloc[i], 0, 0, df.get(column_headers[1]).iloc[i], 0, 0]

    treedata = sg.TreeData()

    for i in range(dfColumnNumber):
        if column_headers[i] != None:
            treedata.Insert('', column_headers[i], column_headers[i], data2003[i], icon=check[0])

def updateTree():

    global data
    global updatedHeadings
    updatedHeadings = [*headings]

    if len(treeSeq) == 0:
        for i in range(3):
            data[i][0] = 0
            data[i][3] = 0
            updatedHeadings[0] = ''
            updatedHeadings[3] = ''
    elif len(treeSeq) == 1:
        for i in range(3):
            data[i][3] = 0
            updatedHeadings[3] = ''
            updatedHeadings[0] = column_headers[treeSeq[len(treeSeq)-1]]
            data[i][0] = df.get(updatedHeadings[0]).iloc[i]
    else:
        for i in range(3):
            updatedHeadings[3] = column_headers[treeSeq[len(treeSeq)-1]]
            updatedHeadings[0] = column_headers[treeSeq[len(treeSeq)-2]]
            data[i][3] = df.get(updatedHeadings[3]).iloc[i]
            data[i][0] = df.get(updatedHeadings[0]).iloc[i]

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def createReport(progress_bar, porgress_percentage):
    if column_headers[0] == '':
        return
    np.set_printoptions(precision=3)
    if outputType != 2:
        workbook = xlsxwriter.Workbook('Output.xlsx')

        sheetSeq = [i + 1 for i in range(len(treeSeq))]
        for n in range(len(treeSeq)):
            
            worksheet = workbook.add_worksheet(str(column_headers[treeSeq[n]]))

            bold_format = workbook.add_format({'bold': True})

            worksheet.write(0, 3, column_headers[treeSeq[n]])
            
            for i in range(1, shotNumber + 1):
                worksheet.write(1, i, "Shot " + str(i), bold_format)
        
            worksheet.write(1, shotNumber + 1, "Max", bold_format)
            worksheet.write(1, shotNumber + 2, "Min", bold_format)

            delta_symbol = u'\u0394'  # This is the unicode for the delta symbol
            worksheet.write(0, 0, delta_symbol)

            worksheet.write(1, shotNumber + 3, delta_symbol + "Prozess", bold_format)

            worksheet.write(1, 0, "Kav.", bold_format)
        
            for i in range(1, kavNumber + 1):
                worksheet.write(i + 1, 0, "Kav " + str(i), bold_format)

            worksheet.write(kavNumber + 2, 0, "max Kav.", bold_format)
            worksheet.write(kavNumber + 3, 0, "min Kav.", bold_format)
            worksheet.write(kavNumber + 4, 0, "Delta Kav.", bold_format)

            list_of_numbers = df.get(column_headers[treeSeq[n]])

            numpy_array = np.array(list_of_numbers)[0:shotNumber * kavNumber]

            reshaped_array = numpy_array.reshape(shotNumber, kavNumber)

            transposed_matrix = np.transpose(reshaped_array)

            max_value_Kav = np.around(np.max(transposed_matrix, axis=0), 3)
            min_value_Kav = np.around(np.min(transposed_matrix, axis=0), 3)
            delta_Kav = max_value_Kav - min_value_Kav
            max_delta_Kav = np.around(np.max(delta_Kav), 3)

            max_value_shot = np.around(np.max(transposed_matrix, axis=1), 3)
            min_value_shot = np.around(np.min(transposed_matrix, axis=1), 3)

            delta_shot = max_value_shot - min_value_shot
            max_delta_shot = np.around(np.max(delta_shot), 3)

            transposed_matrix = np.around(np.transpose(reshaped_array), 4)

            total_delta = max_delta_Kav + max_delta_shot

            for i in range(1, kavNumber + 1):
                for j in range(1, shotNumber + 2):
                    if j < shotNumber + 1:
                        worksheet.write(i + 1, j, transposed_matrix[i - 1][j - 1])
                    else:
                        worksheet.write(i + 1, j, max_value_shot[i - 1])
                        worksheet.write(i + 1, j + 1, min_value_shot[i - 1])
                        worksheet.write(i + 1, j + 2, delta_shot[i - 1])

            for j in range(1, shotNumber + 1):
                i = kavNumber + 1
                worksheet.write(i + 1, j, max_value_Kav[j - 1])
                worksheet.write(i + 2, j, min_value_Kav[j - 1])
                worksheet.write(i + 3, j, delta_Kav[j - 1])
            
            row = kavNumber + 2
            col = shotNumber + 1

            merge_format = workbook.add_format(
                {
                    "bold": 1,
                    "border": 1,
                    "align": "center",
                    "valign": "vcenter",
                    "fg_color": "#d3d3d3",
                }
            )
            # Merge cells.
            a = ord('A') + shotNumber
            y = chr(a + 3) + '1'
            worksheet.merge_range("A1:" + y, "QM" + str(column_headers[treeSeq[n]]) + " = K2002 (MERKMAL)", merge_format)

            x1 = kavNumber + 3
            y1 = chr(a + 2)
            yy1 = chr(a + 3)
            x2 = kavNumber + 4
            x3 = kavNumber + 5

            worksheet.write(row, col, "Nominal", bold_format)
            row += 1
            worksheet.write(row, col, "USL:", bold_format)
            row += 1
            worksheet.write(row, col, "LSL:", bold_format)

            worksheet.merge_range(y1 + str(x1) + ":" + yy1 + str(x1), str(dataOther[treeSeq[n]][0]), merge_format)
            worksheet.merge_range(y1 + str(x2) + ":" + yy1 + str(x2), str(dataOther[treeSeq[n]][1]), merge_format)
            worksheet.merge_range(y1 + str(x3) + ":" + yy1 + str(x3), str(dataOther[treeSeq[n]][2]), merge_format)

            row += 2
            # worksheet.write(row, 2, "max Kav.", bold_format)

            worksheet.merge_range("C" + str(row + 1) + ":" + "D" + str(row + 1), "max Kav.", merge_format)

            worksheet.write(row, 4, np.max(max_value_Kav))
            # worksheet.write(row, 4, "Max Prozess", bold_format)
            worksheet.merge_range("F" + str(row + 1) + ":" + "G" + str(row + 1), "Max Prozess", merge_format)

            worksheet.write(row, 7, np.max(max_value_shot))

            row += 1
            # worksheet.write(row, 2, "min Kav.", bold_format)
            worksheet.merge_range("C" + str(row + 1) + ":" + "D" + str(row + 1), "min Kav.", merge_format)

            worksheet.write(row, 4, np.min(min_value_Kav))
            # worksheet.write(row, 4, "Min Prozess", bold_format)
            worksheet.merge_range("F" + str(row + 1) + ":" + "G" + str(row + 1), "Min Prozess", merge_format)

            worksheet.write(row, 7, np.min(min_value_shot))

            row += 1
            # worksheet.write(row, 2, "Max " + delta_symbol + "Kav", bold_format)
            worksheet.merge_range("C" + str(row + 1) + ":" + "D" + str(row + 1), "Max " + delta_symbol + "Kav", merge_format)
            worksheet.write(row, 4, max_delta_Kav)
            # worksheet.write(row, 4, "Max" + delta_symbol + "Prozess", bold_format)
            worksheet.merge_range("F" + str(row + 1) + ":" + "G" + str(row + 1), "Max" + delta_symbol + "Prozess", merge_format)
            worksheet.write(row, 7, max_delta_shot)

            row += 1
            worksheet.merge_range("C" + str(row + 1) + ":" + "G" + str(row + 1), "Total" + delta_symbol, merge_format)
            worksheet.write(row, 7, total_delta)

            border_format=workbook.add_format({
                                'border':1,
                                'align':'left',
                                'font_size':10
                            })

            worksheet.conditional_format( 'A1:O46' , { 'type' : 'no_blanks' , 'format' : border_format} )

            x = np.arange(1, shotNumber + 1)

            t = np.array([i for i in range(1, shotNumber + 1)])
            nominal = np.full(shotNumber, dataOther[treeSeq[n]][0])
            usl = np.full(shotNumber, dataOther[treeSeq[n]][1])
            lsl = np.full(shotNumber, dataOther[treeSeq[n]][2])

            fig,ax=plt.subplots()
            for i in range(kavNumber):
                ax.scatter(
                    x , transposed_matrix[i], label="Kav " + str(i + 1),
                    s=50, color='black', alpha=0.7
                )

            ax.plot(t, nominal, label = "Nominal", color='black', linestyle='dashed')
            ax.plot(t, usl, label = "USL", color='black', linestyle='solid')
            ax.plot(t, lsl, label = "LSL", color='black', linestyle='solid')

            fig.subplots_adjust(bottom=0.2)

            ax.legend(
                loc="upper center",
                ncol=6,
                bbox_to_anchor=(0.5, 0.16),
                bbox_transform=fig.transFigure 
            )

            imgdata=io.BytesIO()
            fig.savefig(imgdata, format='png')
            worksheet.insert_image(row + 5, 0, '', {'image_data': imgdata})
            
            global progressPercentage
            progressPercentage = int((n + 1) / len(treeSeq) * 100)

            progress_bar.update(current_count=progressPercentage)
            porgress_percentage.update(f'{progressPercentage}%')

        workbook.close()

    if outputType == 1:

        o = win32com.client.Dispatch("Excel.Application")

        o.Visible = False

        current_path = os.getcwd()

        wb = o.Workbooks.Open(current_path+'\\Output.xlsx')

        wb.WorkSheets(sheetSeq).Select()

        wb.ActiveSheet.ExportAsFixedFormat(0, current_path+'\\Output.pdf')

        wb.Close(True)

        os.remove("Output.xlsx")
    
    if outputType == 2:

        document = Document()
        sheetSeq = [i + 1 for i in range(len(treeSeq))]
        for n in range(len(treeSeq)):

            list_of_numbers = df.get(column_headers[treeSeq[n]])

            numpy_array = np.array(list_of_numbers)[0:shotNumber * kavNumber]

            reshaped_array = numpy_array.reshape(shotNumber, kavNumber)

            transposed_matrix = np.transpose(reshaped_array)

            max_value_Kav = np.around(np.max(transposed_matrix, axis=0), 3)
            min_value_Kav = np.around(np.min(transposed_matrix, axis=0), 3)
            delta_Kav = np.around(max_value_Kav - min_value_Kav, 3)
            max_delta_Kav = np.max(delta_Kav)

            max_value_shot = np.around(np.max(transposed_matrix, axis=1), 3)
            min_value_shot = np.around(np.min(transposed_matrix, axis=1), 3)

            delta_shot = np.around(max_value_shot - min_value_shot, 3)
            max_delta_shot = np.around(np.max(delta_shot), 3)

            total_delta = max_delta_Kav + max_delta_shot

            transposed_matrix = np.around(transposed_matrix, 4)

            table = document.add_table(rows=kavNumber+5, cols=shotNumber+4)
            table.style = 'TableGrid'

            for row_height in table.rows:
                row_height.height = Cm(0.5)
            # table.style = 'MediumGrid3'

            delta_symbol = u'\u0394'

            table.cell(0, 0).merge(table.cell(0, shotNumber + 3))

            table.cell(0, 0).text = "QM" + str(str(column_headers[treeSeq[n]]) + " = K2002 (MERKMAL)")
            cell = table.cell(0, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(1, 0).text = "Kav."
            cell = table.cell(1, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(kavNumber + 2, 0).text = "max Kav."
            cell = table.cell(kavNumber + 2, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(kavNumber + 3, 0).text = "min Kav."
            cell = table.cell(kavNumber + 3, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(kavNumber + 4, 0).text = delta_symbol + "Kav."
            cell = table.cell(kavNumber + 4, 0)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True

            row = kavNumber + 2
            col = shotNumber + 1

            table.cell(row, col).merge(table.cell(row, col + 1))
            table.cell(row, col).text = "Nominal"
            cell = table.cell(row, col)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(row, col + 2).text = str(dataOther[treeSeq[n]][0])

            row += 1
            table.cell(row, col).merge(table.cell(row, col + 1))
            table.cell(row, col).text = "USL:"
            cell = table.cell(row, col)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(row, col + 2).text = str(dataOther[treeSeq[n]][1])

            row += 1
            table.cell(row, col).merge(table.cell(row, col + 1))
            table.cell(row, col).text = "LSL:"
            cell = table.cell(row, col)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                     run.bold = True
            table.cell(row, col + 2).text = str(dataOther[treeSeq[n]][2])

            for i in range(2, kavNumber + 2):
                table.cell(i, 0).text = "Kav. " + str(i - 1)
                cell = table.cell(i, 0)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                table.cell(i, shotNumber + 1).text = str(max_value_shot[i - 2])
                table.cell(i, shotNumber + 2).text = str(min_value_shot[i - 2])
                table.cell(i, shotNumber + 3).text = str(delta_shot[i - 2])
                for j in range(1, shotNumber + 1):
                    table.cell(i, j).text = str(transposed_matrix[i - 2][j-1])

            delta_symbol = u'\u0394'
            table.cell(1, shotNumber + 1).text = "Max"
            cell = table.cell(1, shotNumber + 1)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table.cell(1, shotNumber + 2).text = "Min"
            cell = table.cell(1, shotNumber + 2)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table.cell(1, shotNumber + 3).text = delta_symbol + "Prozess"
            cell = table.cell(1, shotNumber + 3)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for i in range(1, shotNumber + 1):
                table.cell(1, i).text = "Shot" + str(i)
                cell = table.cell(1, i)
                for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                table.cell(kavNumber + 2, i).text = str(max_value_Kav[i - 1])
                table.cell(kavNumber + 3, i).text = str(min_value_Kav[i - 1])
                table.cell(kavNumber + 4, i).text = str(delta_Kav[i - 1])

            for r in table.rows:
                for c in r.cells:
                    pp = c.paragraphs
                    for p in pp:
                        # Set the line spacing to center the content vertically
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = Pt(13)  # Adjust the line spacing value as needed to center the content
                        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                        for ru in p.runs:
                            font = ru.font
                            font.size = Pt(6.5)
            
            paragraph = document.add_paragraph('  ')

            table_small = document.add_table(rows=4, cols=4)

            table_small.style = 'TableGrid'

            for row_height in table_small.rows:
                row_height.height = Cm(0.5)

            table_small.cell(0, 0).text = "max Kav."
            cell = table_small.cell(0, 0)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(0, 1).text = str(np.max(max_value_Kav))
            table_small.cell(0, 2).text = "Max Prozess"
            cell = table_small.cell(0, 2)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(0, 3).text = str(np.max(max_value_shot))

            table_small.cell(1, 0).text = "min Kav."
            cell = table_small.cell(1, 0)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(1, 1).text = str(np.min(min_value_Kav))
            table_small.cell(1, 2).text = "min Prozess"
            cell = table_small.cell(1, 2)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(1, 3).text = str(np.min(min_value_shot))

            delta_symbol = u'\u0394'
            table_small.cell(2, 0).text = "Max " + delta_symbol + "Kav"
            cell = table_small.cell(2, 0)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(2, 1).text = str(max_delta_Kav)
            table_small.cell(2, 2).text = "Max " + delta_symbol + "Prozess"
            cell = table_small.cell(2, 2)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(2, 3).text = str(max_delta_shot)

            table_small.cell(3, 0).merge(table.cell(3, 2))
            table_small.cell(3, 0).text = "Total" + delta_symbol
            cell = table_small.cell(3, 0)
            for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            table_small.cell(3, 3).text = str(total_delta)

            for r in table_small.rows:
                for c in r.cells:
                    pp = c.paragraphs
                    for p in pp:
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = Pt(13)  # Adjust the line spacing value as needed to center the content
                        p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                        for ru in p.runs:
                            font = ru.font
                            font.size = Pt(6.5)

            # paragraph = document.add_paragraph('  ')

            x = np.arange(1, shotNumber + 1)
            t = np.array([i for i in range(1, shotNumber + 1)])
            nominal = np.full(shotNumber, dataOther[treeSeq[n]][0])
            usl = np.full(shotNumber, dataOther[treeSeq[n]][1])
            lsl = np.full(shotNumber, dataOther[treeSeq[n]][2])
            
            fig,ax=plt.subplots()
            # fig.suptitle('test title', fontsize=20)
            plt.xlabel('xlabel', fontsize=6.5)
            plt.ylabel('ylabel', fontsize=6.5)
            for i in range(kavNumber):
                ax.scatter(
                    x , transposed_matrix[i], label="Kav " + str(i + 1),
                    s=50, color='black', alpha=0.7
                )
            for label in (ax.get_xticklabels() + ax.get_yticklabels()):
                label.set_fontsize(6.5)  # setting the font size for the tick labels
            ax.plot(t, nominal, label = "Nominal", color='black', linestyle='dashed')
            ax.plot(t, usl, label = "USL", color='black', linestyle='solid')
            ax.plot(t, lsl, label = "LSL", color='black', linestyle='solid')

            fig.subplots_adjust(bottom=0.2)

            ax.legend(
                loc="upper center",
                ncol=6,
                bbox_to_anchor=(0.5, 0.16),
                bbox_transform=fig.transFigure,
                fontsize=6.5
            )

            imgdata=io.BytesIO()
            fig.savefig(imgdata, format='png')

            document.add_picture(imgdata)
            
            document.add_page_break()

            progressPercentage = int((n + 1) / len(treeSeq) * 100)

            progress_bar.update(current_count=progressPercentage)
            porgress_percentage.update(f'{progressPercentage}%')

        document.save('output.docx')
        
def make_window(theme):
    sg.theme(theme)

    col_layout1 = [[sg.Button("1.1 Modify", size=(20, 1))]]

    layout_l = [[sg.Button("1. Load Data", size=(20, 1)),sg.Column(col_layout1, element_justification='right', expand_x=True)],
                [sg.Table(values=data[:][:], headings=headings, 
                  auto_size_columns=False,
                  justification='center',
                  num_rows=10,
                  key='-TABLE-',
                  selected_row_colors='red on yellow',
                  enable_events=True,
                  expand_x=True,
                  expand_y=True,
                  vertical_scroll_only=False,  # Remove horizontal scroll bar
                  tooltip='This is a MERKMAL data')]]
    
    col_layout2 = [[sg.Button("3. Create Report", size=(20, 1))]]
    
    layout_r = [[sg.Button("2. Select Merkmal", size=(20, 1))],
               
        [sg.Tree(data=treedata, headings=headingsTree[1:], auto_size_columns=True,
        num_rows=7, col0_width=20, key='-TREE-', row_height=24, metadata=[],
        show_expanded=False, enable_events=True,
        select_mode=sg.TABLE_SELECT_MODE_BROWSE)],
        [sg.Radio('All', "RadioDemo", size=(10,1), key='-All-',enable_events=True), sg.Radio('None', "RadioDemo", size=(10,1), key='-None-',enable_events=True)]]

    layout_rr = [[sg.Column(col_layout2, element_justification='right', expand_x=True) ],
             [sg.Combo(choices, default_value='.xlsx', readonly=True,enable_events=True, key='-COMBO-')]

             ]

    layout = [
         [sg.Column(layout_l, size=(600, 260)), sg.Column(layout_r, size=(390, 260)),sg.Column(layout_rr, size=(190, 260))],
         [sg.ProgressBar(100, orientation='h', size=(50, 20), key='-PROGRESS BAR-'), sg.Text(f'{progressPercentage}%', size=(6, 1), key=('-%-'))]
        ]
    window = sg.Window('AQDEF Reader', layout, grab_anywhere=True, resizable=True, margins=(0, 0),
                       use_custom_titlebar=True, finalize=True, keep_on_top=True)
    window['-TREE-'].Widget.heading("#0", text=headingsTree[0])
    return window

def main():
    window = make_window(sg.theme())
    while True:
        event, values = window.read(timeout=100)
        if event in (None, 'Exit'):
            print("[LOG] Clicked Exit!")
            break
        elif event == "1. Load Data":
            window['-All-'].update(value=False)
            window['-None-'].update(value=False)
            treeSeq.clear()
            print("[LOG] Clicked Load File!")
            folder_or_file = sg.popup_get_file('Choose your file', keep_on_top=True)
            
            if folder_or_file:
                processDFQ(folder_or_file)
                window['-TABLE-'].update(values=data)
                window['-TREE-'].update(values=treedata)
        elif event == "1.1 Modify":
            print("[LOG] Clicked Modify!")
            SetNumber()
            window['-TABLE-'].update(values=data)
        elif event == "3. Create Report":
            print("[LOG] Created Report!")
            progress_bar = window['-PROGRESS BAR-']
            porgress_percentage = window[('-%-')]
            createReport(progress_bar, porgress_percentage)
        elif event == '-All-':
            window['-TREE-'].metadata.clear()
            treeSeq.clear()
            if column_headers[0] != '':
                for i in range(dfColumnNumber):
                    if column_headers[i] != None:
                        window['-TREE-'].metadata.append(column_headers[i])
                        window['-TREE-'].update(key=column_headers[i], icon=check[1])
                        treeSeq.append(i)
            updateTree()
            window['-TABLE-'].update(values=data)
            update_title(window['-TABLE-'].Widget, updatedHeadings)
        elif event == '-None-':
            window['-TREE-'].metadata.clear()
            treeSeq.clear()
            if column_headers[0] != '':
                for i in range(dfColumnNumber):
                    window['-TREE-'].update(key=column_headers[i], icon=check[0])
            updateTree()
            window['-TABLE-'].update(values=data)
            update_title(window['-TABLE-'].Widget, updatedHeadings)
        elif event == '-TREE-':
            window['-All-'].update(value=False)
            window['-None-'].update(value=False)
            if len(values['-TREE-']) == 0:
                continue 
            president = values['-TREE-'][0]

            if president in window['-TREE-'].metadata:
                window['-TREE-'].metadata.remove(president)
                treeSeq.remove(column_headers.index(president))
                window['-TREE-'].update(key=president, icon=check[0])
            else:
                window['-TREE-'].metadata.append(president)
                treeSeq.append(column_headers.index(president))
                window['-TREE-'].update(key=president, icon=check[1])
            updateTree()
            window['-TABLE-'].update(values=data)
            update_title(window['-TABLE-'].Widget, updatedHeadings)
        elif event == '-COMBO-':
            global outputType
            outputType = choices.index(values['-COMBO-'])
            
    window.close()
    exit(0)

if __name__ == '__main__':
    sg.theme('DefaultNoMoreNagging')
    main()